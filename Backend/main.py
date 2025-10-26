"""
ADE Insurance Document Intelligence Backend API
FastAPI server with mock AI processing capabilities
"""

from fastapi import FastAPI, HTTPException, File, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from contextlib import asynccontextmanager
import os
import uuid
import aiofiles
import time
from datetime import datetime
from typing import Optional, List
import json

# PDF and Image processing imports
try:
    import fitz  # PyMuPDF
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    print("Warning: PyMuPDF not installed. PDF processing will be limited.")

try:
    from PIL import Image
    IMAGE_PROCESSING_AVAILABLE = True
except ImportError:
    IMAGE_PROCESSING_AVAILABLE = False
    print("Warning: Pillow not installed. Image processing will be limited.")

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    DOCX_PROCESSING_AVAILABLE = True
except ImportError:
    DOCX_PROCESSING_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX processing will be limited.")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    print("Warning: pdf2image not installed. Alternative PDF processing will be used.")

from app.database import init_db, get_db
from app.models import Document, Job, Page, User
from app.schemas import (
    DocumentResponse,
    JobResponse,
    UploadResponse,
    RegionResponse,
    ProcessingRequest,
    UserRegister,
    UserLogin,
    UserResponse,
    TokenResponse
)
from app.ai_service import analyze_auto_document, extract_markdown_content, get_image_path_from_url

# JWT Configuration
from datetime import timedelta
import jwt
from typing import Optional

SECRET_KEY = "your-secret-key-change-this-in-production"  # Change this!
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 7 days

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    """Create JWT access token"""
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

# Utility Functions
def merge_page_results(page_results: List[dict]) -> dict:
    """
    Merge analysis results from multiple pages into a single result
    
    Args:
        page_results: List of analysis results from each page
        
    Returns:
        Merged result with combined information from all pages
    """
    if not page_results:
        return {
            "error": "No pages analyzed",
            "document_type": "Error",
            "confidence": 0.0,
            "title": None,
            "summary": "No content available",
            "people": [],
            "organizations": [],
            "locations": [],
            "dates": [],
            "numbers": [],
            "signature_detected": False,
            "total_pages": 0,
            "pages": []
        }
    
    # If only one page, return it with minor modifications
    if len(page_results) == 1:
        result = page_results[0].copy()
        result['total_pages'] = 1
        result['pages'] = [result.copy()]
        return result
    
    # For multiple pages, merge information
    # Collect all unique values across pages
    all_people = []
    all_organizations = []
    all_locations = []
    all_dates = []
    all_numbers = []
    has_signature = False
    
    # Document type determination (most common, highest confidence)
    doc_types = {}
    total_confidence = 0
    titles = []
    summaries = []
    
    for page_result in page_results:
        # Collect people (can be strings or dicts)
        if 'people' in page_result and page_result['people']:
            for person in page_result['people']:
                if isinstance(person, dict):
                    # If dict, try to get name field
                    person_str = person.get('name') or person.get('label') or str(person)
                    if person_str and person_str not in [p if isinstance(p, str) else p.get('name', '') for p in all_people]:
                        all_people.append(person)
                elif isinstance(person, str) and person not in [p if isinstance(p, str) else p.get('name', '') for p in all_people]:
                    all_people.append(person)
        
        # Collect organizations (can be strings or dicts)
        if 'organizations' in page_result and page_result['organizations']:
            for org in page_result['organizations']:
                if isinstance(org, dict):
                    org_str = org.get('name') or org.get('label') or str(org)
                    if org_str and org_str not in [o if isinstance(o, str) else o.get('name', '') for o in all_organizations]:
                        all_organizations.append(org)
                elif isinstance(org, str) and org not in [o if isinstance(o, str) else o.get('name', '') for o in all_organizations]:
                    all_organizations.append(org)
        
        # Collect locations
        if 'locations' in page_result and page_result['locations']:
            for loc in page_result['locations']:
                if isinstance(loc, str) and loc not in all_locations:
                    all_locations.append(loc)
                elif isinstance(loc, dict):
                    loc_str = loc.get('name') or loc.get('label') or str(loc)
                    if loc_str not in all_locations:
                        all_locations.append(loc_str)
        
        # Collect dates
        if 'dates' in page_result and page_result['dates']:
            for date in page_result['dates']:
                if isinstance(date, dict):
                    date_str = date.get('value') or date.get('label') or str(date)
                    if date_str and date_str not in [d if isinstance(d, str) else d.get('value', '') for d in all_dates]:
                        all_dates.append(date)
                elif isinstance(date, str) and date not in [d if isinstance(d, str) else d.get('value', '') for d in all_dates]:
                    all_dates.append(date)
        
        # Collect numbers
        if 'numbers' in page_result and page_result['numbers']:
            for num in page_result['numbers']:
                if isinstance(num, dict):
                    num_str = num.get('value') or num.get('label') or str(num)
                    if num_str and num_str not in [n if isinstance(n, str) else n.get('value', '') for n in all_numbers]:
                        all_numbers.append(num)
                elif isinstance(num, str) and num not in [n if isinstance(n, str) else n.get('value', '') for n in all_numbers]:
                    all_numbers.append(num)
        
        # Check for signatures
        if page_result.get('signature_detected', False):
            has_signature = True
        
        # Count document types
        doc_type = page_result.get('document_type', 'Unknown')
        confidence = page_result.get('confidence', 0.0)
        
        if doc_type not in doc_types:
            doc_types[doc_type] = {'count': 0, 'total_confidence': 0}
        doc_types[doc_type]['count'] += 1
        doc_types[doc_type]['total_confidence'] += confidence
        total_confidence += confidence
        
        # Collect titles and summaries
        if page_result.get('title'):
            titles.append(f"Page {page_result.get('page_number', '?')}: {page_result['title']}")
        if page_result.get('summary'):
            summaries.append(f"Page {page_result.get('page_number', '?')}: {page_result['summary']}")
    
    # Determine most likely document type (weighted by confidence)
    best_doc_type = "Unknown"
    best_score = 0
    for doc_type, info in doc_types.items():
        # Score = count * average confidence
        avg_confidence = info['total_confidence'] / info['count']
        score = info['count'] * avg_confidence
        if score > best_score:
            best_score = score
            best_doc_type = doc_type
    
    # Calculate average confidence
    avg_confidence = total_confidence / len(page_results) if page_results else 0.0
    
    # Create merged title
    merged_title = titles[0] if titles else None
    
    # Create merged summary
    merged_summary = f"Multi-page document with {len(page_results)} pages. " + " | ".join(summaries[:3])
    if len(summaries) > 3:
        merged_summary += f" ... and {len(summaries) - 3} more pages"
    
    # Return merged result
    return {
        "document_type": best_doc_type,
        "confidence": round(avg_confidence, 2),
        "title": merged_title,
        "summary": merged_summary,
        "people": all_people,
        "organizations": all_organizations,
        "locations": all_locations,
        "dates": all_dates,
        "numbers": all_numbers,
        "signature_detected": has_signature,
        "total_pages": len(page_results),
        "pages": page_results  # Include individual page results for reference
    }

# PDF Processing Functions
async def process_pdf_to_images(pdf_path: str, document_id: str) -> List[str]:
    """
    Convert PDF pages to individual images
    Returns list of image URLs (relative to /data/)
    """
    if not PDF_PROCESSING_AVAILABLE:
        # Fallback: return original PDF path
        return [f"/data/docs/{document_id}.pdf"]
    
    image_urls = []
    
    try:
        # Open PDF
        pdf_doc = fitz.open(pdf_path)
        
        for page_num in range(len(pdf_doc)):
            # Get page
            page = pdf_doc.load_page(page_num)
            
            # Convert to image (150 DPI for good quality and reasonable file size)
            mat = fitz.Matrix(1.5, 1.5)  # 1.5x zoom = ~150 DPI
            pix = page.get_pixmap(matrix=mat)
            
            # Save as PNG in images folder
            image_filename = f"{document_id}_page_{page_num + 1}.png"
            image_path = f"data/images/{image_filename}"
            pix.save(image_path)
            
            # Return URL path
            image_urls.append(f"/data/images/{image_filename}")
            
        pdf_doc.close()
        
        # If no pages were processed, return fallback
        if not image_urls:
            return [f"/data/docs/{document_id}.pdf"]
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        # Fallback: return original PDF path
        return [f"/data/docs/{document_id}.pdf"]
    
    return image_urls

async def process_image_file(image_path: str, document_id: str, page_index: int = 0) -> str:
    """
    Process single image file (basic validation)
    Returns image URL path
    """
    # For now, just return the URL path without processing
    # In the future, could add resize/optimization here
    file_ext = os.path.splitext(image_path)[1]
    return f"/data/images/{document_id}{file_ext}"

async def process_docx_to_images(docx_path: str, document_id: str) -> List[str]:
    """
    Convert DOCX pages to individual images
    Returns list of image URLs (relative to /data/)
    """
    if not DOCX_PROCESSING_AVAILABLE:
        # Fallback: return original DOCX path
        return [f"/data/docs/{document_id}.docx"]
    
    image_urls = []
    
    try:
        # For DOCX, we'll create a preview image of the first page
        # This is a simplified approach - in production you might want to use more sophisticated methods
        
        # Load DOCX document
        doc = DocxDocument(docx_path)
        
        # Create a simple preview by extracting text and creating an image
        # For now, we'll just create one preview image
        preview_filename = f"{document_id}_preview.png"
        preview_path = f"data/images/{preview_filename}"
        
        # Create a simple text-based preview image
        if IMAGE_PROCESSING_AVAILABLE:
            from PIL import Image, ImageDraw, ImageFont
            
            # Get document text (first few paragraphs)
            text_content = ""
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() and paragraph_count < 10:  # First 10 paragraphs
                    text_content += paragraph.text.strip() + "\n\n"
                    paragraph_count += 1
            
            if not text_content:
                text_content = "DOCX Document Preview\n\nDocument contains formatted content."
            
            # Create image with text
            img_width, img_height = 800, 1000
            img = Image.new('RGB', (img_width, img_height), color='white')
            draw = ImageDraw.Draw(img)
            
            try:
                # Try to use a system font
                font = ImageFont.load_default()
            except:
                font = None
            
            # Draw text with word wrapping
            y_position = 50
            lines = text_content.split('\n')
            for line in lines[:30]:  # Limit to 30 lines
                if y_position > img_height - 50:
                    break
                draw.text((50, y_position), line[:80], fill='black', font=font)  # Limit line length
                y_position += 25
            
            # Save preview image
            img.save(preview_path, "PNG")
            image_urls.append(f"/data/images/{preview_filename}")
        else:
            # If PIL not available, return DOCX file path
            return [f"/data/docs/{document_id}.docx"]
        
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        # Fallback: return original DOCX path
        return [f"/data/docs/{document_id}.docx"]
    
    return image_urls if image_urls else [f"/data/docs/{document_id}.docx"]

async def process_pdf_alternative(pdf_path: str, document_id: str) -> List[str]:
    """
    Alternative PDF processing using pdf2image
    Returns list of image URLs (relative to /data/)
    """
    if not PDF2IMAGE_AVAILABLE:
        return await process_pdf_to_images(pdf_path, document_id)
    
    image_urls = []
    
    try:
        # Convert PDF pages to images using pdf2image
        images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=10)  # Limit to first 10 pages
        
        for i, image in enumerate(images):
            image_filename = f"{document_id}_page_{i + 1}.png"
            image_path = f"data/images/{image_filename}"
            
            # Save image
            image.save(image_path, "PNG")
            image_urls.append(f"/data/images/{image_filename}")
            
        # If no pages were processed, fallback
        if not image_urls:
            return [f"/data/docs/{document_id}.pdf"]
        
    except Exception as e:
        print(f"Error in alternative PDF processing: {e}")
        # Fallback to original PDF processing
        return await process_pdf_to_images(pdf_path, document_id)
    
    return image_urls

@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    await init_db()
    # Create data directories if they don't exist
    os.makedirs("data/docs", exist_ok=True)
    os.makedirs("data/images", exist_ok=True)
    yield
    # Shutdown (cleanup if needed)

# Initialize FastAPI app
app = FastAPI(
    title="ADE Insurance Document Intelligence API",
    description="Mock AI Document Processing API for Insurance Documents",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc",
    lifespan=lifespan
)

# Configure CORS for Frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:5174", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files (uploaded documents and images)
app.mount("/data", StaticFiles(directory="data"), name="data")

@app.get("/")
async def root():
    """Fast health check endpoint"""
    return {
        "status": "healthy",
        "message": "ADE Insurance Document Intelligence API",
        "version": "1.0.0",
        "docs": "/docs"
    }

@app.get("/health")
async def health_check():
    """Fast health check for monitoring"""
    return {"status": "healthy", "timestamp": time.time()}

# ==================== Authentication Endpoints ====================

@app.post("/auth/register", response_model=TokenResponse)
async def register_user(user_data: UserRegister):
    """
    Register a new user account
    """
    db = get_db()
    try:
        # Check if email already exists
        existing_user = db.query(User).filter(User.email == user_data.email).first()
        if existing_user:
            raise HTTPException(
                status_code=400,
                detail="Email already registered"
            )
        
        # Create new user
        hashed_password = User.hash_password(user_data.password)
        new_user = User(
            email=user_data.email,
            full_name=user_data.full_name,
            phone=user_data.phone,
            hashed_password=hashed_password
        )
        
        db.add(new_user)
        db.commit()
        db.refresh(new_user)
        
        # Create access token
        access_token = create_access_token(
            data={"sub": str(new_user.id), "email": new_user.email},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        
        # Return token and user data
        return TokenResponse(
            access_token=access_token,
            user=UserResponse.model_validate(new_user)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=500,
            detail=f"Registration failed: {str(e)}"
        )
    finally:
        db.close()

@app.post("/auth/login", response_model=TokenResponse)
async def login_user(credentials: UserLogin):
    """
    Login user and return JWT token
    """
    db = get_db()
    try:
        # Find user by email
        user = db.query(User).filter(User.email == credentials.email).first()
        
        if not user or not user.verify_password(credentials.password):
            raise HTTPException(
                status_code=401,
                detail="Incorrect email or password"
            )
        
        if not user.is_active:
            raise HTTPException(
                status_code=403,
                detail="Account is inactive"
            )
        
        # Create access token
        access_token = create_access_token(
            data={"sub": str(user.id), "email": user.email},
            expires_delta=timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        )
        
        # Return token and user data
        return TokenResponse(
            access_token=access_token,
            user=UserResponse.model_validate(user)
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Login failed: {str(e)}"
        )
    finally:
        db.close()

@app.get("/auth/me", response_model=UserResponse)
async def get_current_user(token: str):
    """
    Get current user from JWT token
    """
    db = get_db()
    try:
        # Decode JWT token
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        user_id = int(payload.get("sub"))
        
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        # Get user from database
        user = db.query(User).filter(User.id == user_id).first()
        
        if not user:
            raise HTTPException(status_code=404, detail="User not found")
        
        return UserResponse.model_validate(user)
        
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()

# ==================== Document Endpoints ====================

@app.post("/documents/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload a document file and create a new document record
    """
    db = get_db()
    try:
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Get file extension
        file_ext = os.path.splitext(file.filename)[1].lower()
        if file_ext not in ['.pdf', '.docx', '.png', '.jpg', '.jpeg']:
            raise HTTPException(status_code=400, detail="Unsupported file format. Only PDF, DOCX, PNG, JPG are allowed.")
        
        print(f"Uploading document: {file.filename} (ID: {document_id})")
        
        # Save file to appropriate storage location
        if file_ext in ['.pdf', '.docx']:
            file_path = f"data/docs/{document_id}{file_ext}"
        else:
            file_path = f"data/images/{document_id}{file_ext}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
        
        print(f"File saved to: {file_path}")
        
        # Create document record in database
        document = Document(
            id=document_id,
            filename=file.filename,
            status="NOT_STARTED",
            created_at=datetime.utcnow()
        )
        db.add(document)
        
        print(f"Creating document record in database...")
        
        # Create page records based on file type
        if file_ext == '.pdf':
            # For PDF: try to extract pages, fallback to single page if failed
            try:
                print(f"Processing PDF...")
                # Try alternative PDF processing first if available
                if PDF2IMAGE_AVAILABLE:
                    image_urls = await process_pdf_alternative(file_path, document_id)
                else:
                    image_urls = await process_pdf_to_images(file_path, document_id)
                
                print(f"PDF processed into {len(image_urls)} pages")
                
                # Create page for each extracted image/page
                for i, image_url in enumerate(image_urls):
                    page = Page(
                        document_id=document_id,
                        page_index=i,
                        image_url=image_url
                    )
                    db.add(page)
            except Exception as e:
                print(f"PDF processing failed: {e}")
                # Fallback: single page pointing to PDF
                page = Page(
                    document_id=document_id,
                    page_index=0,
                    image_url=f"/data/docs/{document_id}{file_ext}"
                )
                db.add(page)
        elif file_ext == '.docx':
            # For DOCX: try to create preview images
            try:
                print(f"Processing DOCX...")
                image_urls = await process_docx_to_images(file_path, document_id)
                print(f"DOCX processed into {len(image_urls)} pages")
                # Create page for each preview image
                for i, image_url in enumerate(image_urls):
                    page = Page(
                        document_id=document_id,
                        page_index=i,
                        image_url=image_url
                    )
                    db.add(page)
            except Exception as e:
                print(f"DOCX processing failed: {e}")
                # Fallback: single page pointing to DOCX
                page = Page(
                    document_id=document_id,
                    page_index=0,
                    image_url=f"/data/docs/{document_id}{file_ext}"
                )
                db.add(page)
        else:
            # For images: single page pointing to the image
            print(f"Processing image file...")
            page = Page(
                document_id=document_id,
                page_index=0,
                image_url=f"/data/images/{document_id}{file_ext}"
            )
            db.add(page)
        
        db.commit()
        print(f"✅ Document uploaded successfully: {document_id}")
        
        return UploadResponse(document_id=document_id)
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Upload error: {e}")
        import traceback
        traceback.print_exc()
        db.rollback()
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")
    finally:
        db.close()

@app.post("/documents/{document_id}/process")
async def process_document(document_id: str):
    """
    Process document using Gemini auto-analysis
    This replaces the old mock processing system
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages:
            raise HTTPException(status_code=404, detail="No pages found for this document")
        
        # Get first page image URL
        first_page = pages[0]
        image_url = first_page.image_url
        
        # Convert URL to local file path
        image_path = get_image_path_from_url(image_url)
        
        if not image_path:
            raise HTTPException(status_code=400, detail="Invalid image path")
        
        try:
            # Update document status to processing
            document.status = "PROCESSING"
            db.commit()
            
            # Analyze document with Gemini
            result = await analyze_auto_document(image_path)
            
            # Save result to database
            document.ai_result_json = json.dumps(result, ensure_ascii=False, indent=2)
            document.status = "DONE"
            db.commit()
            
            return {
                "status": "DONE",
                "message": "Document processed successfully"
            }
            
        except Exception as e:
            # Update status to error
            document.status = "ERROR"
            db.commit()
            
            raise HTTPException(
                status_code=500, 
                detail=f"Processing failed: {str(e)}"
            )
    finally:
        db.close()

@app.get("/documents/{document_id}", response_model=DocumentResponse)
async def get_document(document_id: str):
    """
    Get document metadata and pages
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get document pages
        pages = db.query(Page).filter(Page.document_id == document_id).all()
        
        pages_data = [
            {
                "page_index": page.page_index,
                "image_url": page.image_url
            }
            for page in pages
        ]
        
        return DocumentResponse(
            document_id=document.id,
            status=document.status,
            pages=pages_data
        )
    finally:
        db.close()

@app.get("/documents/{document_id}/overlay")
async def get_document_overlay(document_id: str):
    """
    Get document overlay regions - removed (not needed for fast mode)
    """
    raise HTTPException(
        status_code=404, 
        detail="Overlay feature not available in fast analysis mode"
    )

@app.get("/documents/{document_id}/markdown")
async def get_document_markdown(document_id: str):
    """
    Get document content as markdown (extracted from Gemini)
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Return markdown if available
        if document.markdown_content:
            return {"markdown": document.markdown_content}
        
        # If not available, check document status
        if document.status == "PROCESSING":
            raise HTTPException(status_code=400, detail="Document is being processed")
        elif document.status == "ERROR":
            raise HTTPException(status_code=400, detail="Document processing failed")
        else:
            raise HTTPException(status_code=400, detail="Document not yet analyzed. Please call /analyze-auto first")
    finally:
        db.close()

@app.get("/documents/{document_id}/json")
async def get_document_json(document_id: str):
    """
    Get document structured data as JSON from Gemini analysis
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Check if document has been analyzed
        if not document.ai_result_json:
            if document.status == "PROCESSING":
                raise HTTPException(status_code=400, detail="Document is being processed")
            elif document.status == "ERROR":
                raise HTTPException(status_code=400, detail="Document processing failed")
            else:
                raise HTTPException(status_code=400, detail="Document not yet analyzed. Please call /analyze-auto first")
        
        # Return AI analysis result
        try:
            return json.loads(document.ai_result_json)
        except json.JSONDecodeError:
            raise HTTPException(status_code=500, detail="Invalid JSON data stored")
    finally:
        db.close()

@app.put("/documents/{document_id}/json")
async def update_document_json(document_id: str, json_data: dict):
    """
    Update document JSON data
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # In a real implementation, you would save this to database
        # For now, just return success
        return {"success": True, "message": "JSON data updated successfully"}
    finally:
        db.close()

@app.post("/documents/{document_id}/analyze-auto")
async def analyze_document_auto(document_id: str):
    """
    Analyze document automatically using Gemini 2.5 Flash
    For multi-page PDFs: analyzes each page separately and merges results
    Extracts structured information and full text markdown
    """
    db = get_db()
    try:
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        # Get ALL pages for this document
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages:
            raise HTTPException(status_code=404, detail="No pages found for this document")
        
        try:
            # Update document status to processing
            document.status = "PROCESSING"
            db.commit()
            
            print(f"\n{'='*80}")
            print(f"📄 Analyzing document: {document_id}")
            print(f"   Filename: {document.filename}")
            print(f"   Total pages: {len(pages)}")
            print(f"{'='*80}\n")
            
            # Lists to collect results from all pages
            all_page_results = []
            all_markdown_parts = []
            
            # Process each page
            for idx, page in enumerate(pages):
                page_num = idx + 1
                print(f"\n📑 Processing page {page_num}/{len(pages)}...")
                print(f"   Image URL: {page.image_url}")
                
                # Convert URL to local file path
                image_path = get_image_path_from_url(page.image_url)
                
                if not image_path:
                    print(f"   ⚠️  Warning: Invalid image path for page {page_num}, skipping...")
                    continue
                
                if not os.path.exists(image_path):
                    print(f"   ⚠️  Warning: Image file not found: {image_path}, skipping...")
                    continue
                
                try:
                    # Analyze this page with Gemini (structured data)
                    print(f"   🤖 Extracting structured data...")
                    page_result = await analyze_auto_document(image_path)
                    
                    # Add page number to result
                    page_result['page_number'] = page_num
                    all_page_results.append(page_result)
                    print(f"   ✅ Structured data extracted")
                    
                    # Extract markdown content (full text)
                    print(f"   📝 Extracting markdown content...")
                    page_markdown = await extract_markdown_content(image_path)
                    
                    # Add page separator and page number to markdown
                    if len(pages) > 1:
                        markdown_with_header = f"\n\n---\n## Page {page_num}\n\n{page_markdown}"
                    else:
                        markdown_with_header = page_markdown
                    
                    all_markdown_parts.append(markdown_with_header)
                    print(f"   ✅ Markdown extracted ({len(page_markdown)} chars)")
                    
                except Exception as page_error:
                    print(f"   ❌ Error analyzing page {page_num}: {page_error}")
                    # Continue with next page even if this one fails
                    all_page_results.append({
                        "page_number": page_num,
                        "error": str(page_error),
                        "document_type": "Error",
                        "confidence": 0.0,
                        "title": None,
                        "summary": f"Page {page_num} analysis failed",
                        "people": [],
                        "organizations": [],
                        "locations": [],
                        "dates": [],
                        "numbers": [],
                        "signature_detected": False
                    })
                    all_markdown_parts.append(f"\n\n---\n## Page {page_num}\n\n*Error extracting content from this page*\n")
            
            # Merge results from all pages
            print(f"\n📊 Merging results from {len(all_page_results)} pages...")
            merged_result = merge_page_results(all_page_results)
            
            # Combine all markdown parts
            full_markdown = "\n".join(all_markdown_parts).strip()
            
            # Add document summary at the top of markdown
            if len(pages) > 1:
                markdown_header = f"# {document.filename}\n\n**Total Pages:** {len(pages)}\n**Document Type:** {merged_result.get('document_type', 'Unknown')}\n"
                full_markdown = markdown_header + full_markdown
            
            print(f"   ✅ Merged result - Document type: {merged_result.get('document_type')}")
            print(f"   ✅ Total markdown length: {len(full_markdown)} chars")
            
            # Save results to database
            document.ai_result_json = json.dumps(merged_result, ensure_ascii=False, indent=2)
            document.markdown_content = full_markdown
            document.status = "DONE"
            db.commit()
            
            print(f"\n✅ Analysis complete for document {document_id}")
            print(f"{'='*80}\n")
            
            return merged_result
            
        except Exception as e:
            # Update status to error
            document.status = "ERROR"
            db.commit()
            
            print(f"\n❌ Analysis failed: {e}")
            import traceback
            traceback.print_exc()
            
            raise HTTPException(
                status_code=500, 
                detail=f"Analysis failed: {str(e)}"
            )
    finally:
        db.close()


@app.post("/documents/{document_id}/extract-person-info")
async def extract_person_info_endpoint(document_id: str):
    """
    Extract personal information from document (CCCD/ID/Driver License)
    Optimized for insurance application forms
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n👤 Extracting person info from document {document_id}")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, extract_person_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Extract person info using Gemini
        person_info = await extract_person_info(image_path)
        
        if "error" in person_info:
            print(f"   ❌ Extraction error: {person_info['error']}")
        else:
            print(f"   ✅ Extracted: {person_info.get('fullName', 'N/A')}")
            
            # Save person_info to database
            import json
            document.person_data = json.dumps(person_info, ensure_ascii=False)
            db.commit()
            print(f"   💾 Saved person data to database")
        
        return {
            "document_id": document_id,
            "person_info": person_info,
            "message": "Person information extracted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Person info extraction failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")
    finally:
        db.close()


@app.post("/documents/{document_id}/extract-vehicle-info")
async def extract_vehicle_info_endpoint(document_id: str):
    """
    Extract vehicle information from document (Giấy đăng ký xe / Cà vẹt)
    For vehicle insurance applications
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n🚗 Extracting vehicle info from document {document_id}")
        
        # Get first page image
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, extract_vehicle_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Extract vehicle info using Gemini
        vehicle_info = await extract_vehicle_info(image_path)
        
        if "error" in vehicle_info:
            print(f"   ❌ Extraction error: {vehicle_info['error']}")
        else:
            print(f"   ✅ Extracted: {vehicle_info.get('licensePlate', 'N/A')} - {vehicle_info.get('brand', 'N/A')} {vehicle_info.get('model', 'N/A')}")
            
            # Save vehicle_info to database
            import json
            document.vehicle_data = json.dumps(vehicle_info, ensure_ascii=False)
            db.commit()
            print(f"   💾 Saved vehicle data to database")
        
        return {
            "document_id": document_id,
            "vehicle_info": vehicle_info,
            "message": "Vehicle information extracted successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Vehicle info extraction failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Extraction failed: {str(e)}")
    finally:
        db.close()


@app.post("/documents/{document_id}/recommend-insurance")
async def recommend_insurance_endpoint(document_id: str):
    """
    Analyze document address and recommend insurance packages based on region
    Returns region-specific insurance recommendations (Bắc/Trung/Nam)
    Uses extracted PersonInfo (placeOfOrigin) if available, otherwise analyzes image
    """
    db = get_db()
    try:
        # Get document
        document = db.query(Document).filter(Document.id == document_id).first()
        if not document:
            raise HTTPException(status_code=404, detail="Document not found")
        
        print(f"\n🏠 Analyzing address for insurance recommendations: {document_id}")
        
        # Get first page
        pages = db.query(Page).filter(Page.document_id == document_id).order_by(Page.page_index).all()
        
        if not pages or not pages[0].image_url:
            raise HTTPException(status_code=400, detail="No image found for this document")
        
        # Get image path
        from app.ai_service import get_image_path_from_url, recommend_insurance_by_address, recommend_insurance_by_person_info
        image_path = get_image_path_from_url(pages[0].image_url)
        
        if not image_path or not os.path.exists(image_path):
            raise HTTPException(status_code=400, detail=f"Image file not found: {image_path}")
        
        print(f"   📷 Processing image: {image_path}")
        
        # Check if we have extracted person info with placeOfOrigin
        person_data = None
        if document.person_data:
            import json
            person_data = json.loads(document.person_data)
        
        if person_data and person_data.get('placeOfOrigin'):
            # Use extracted person info
            print(f"   ✅ Using extracted placeOfOrigin: {person_data.get('placeOfOrigin')}")
            print(f"   ✅ Using extracted address: {person_data.get('address', 'N/A')}")
            recommendation = await recommend_insurance_by_person_info(person_data)
        else:
            # Fallback to image analysis
            print(f"   ⚠️  No person info found, analyzing image directly")
            recommendation = await recommend_insurance_by_address(image_path)
        
        if "error" in recommendation:
            print(f"   ❌ Analysis error: {recommendation['error']}")
        else:
            place_region = recommendation.get('place_of_origin', {}).get('region', 'Unknown')
            addr_region = recommendation.get('address', {}).get('region', 'Unknown')
            num_packages = len(recommendation.get('recommended_packages', []))
            print(f"   ✅ Quê quán region: {place_region}, Address region: {addr_region}")
            print(f"   📦 Recommended: {num_packages} packages")
        
        return {
            "document_id": document_id,
            "recommendation": recommendation,
            "message": "Insurance recommendations generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Insurance recommendation failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Recommendation failed: {str(e)}")
    finally:
        db.close()


@app.post("/chat")
async def chat_endpoint(request: dict):
    """
    Chat with AI Insurance Advisor
    
    Request body:
    {
        "message": "User's message",
        "document_id": "optional document ID for context",
        "chat_history": [optional previous messages]
    }
    """
    db = get_db()
    try:
        message = request.get('message')
        document_id = request.get('document_id')
        chat_history = request.get('chat_history', [])
        
        if not message:
            raise HTTPException(status_code=400, detail="Message is required")
        
        print(f"\n💬 Chat request: message='{message[:50]}...', doc_id={document_id}")
        
        # Get document analysis if document_id provided
        document_analysis = None
        if document_id:
            document = db.query(Document).filter(Document.id == document_id).first()
            if document and document.ai_result_json:
                import json
                ai_result = json.loads(document.ai_result_json)
                
                # Extract recommendation data
                if 'recommendation' in ai_result:
                    document_analysis = ai_result['recommendation']
                    print(f"   📋 Using document analysis with region: {document_analysis.get('place_of_origin', {}).get('region', 'Unknown')}")
        
        # Call chat service
        from app.chat_service import chat_with_insurance_advisor
        response = await chat_with_insurance_advisor(
            user_message=message,
            document_analysis=document_analysis,
            chat_history=chat_history
        )
        
        return {
            "reply": response['reply'],
            "has_context": response.get('has_context', False),
            "region": response.get('region'),
            "message": "Chat response generated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"\n❌ Chat failed: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Chat failed: {str(e)}")
    finally:
        db.close()


# ===================== INSURANCE PURCHASE HISTORY ENDPOINTS =====================

@app.post("/insurance-purchases")
async def create_insurance_purchase(request: dict):
    """
    Create new insurance purchase record
    
    Request body:
    {
        "user_id": 1,
        "package_name": "Bảo hiểm TNDS Xe máy",
        "package_type": "TNDS",
        "customer_name": "Nguyễn Văn A",
        "customer_phone": "0901234567",
        "premium_amount": "500000",
        "document_id": "optional",
        ... (other fields)
    }
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        # Validate required fields
        required_fields = ['user_id', 'package_name', 'package_type', 'customer_name', 'customer_phone', 'premium_amount']
        for field in required_fields:
            if field not in request or not request[field]:
                raise HTTPException(status_code=400, detail=f"Field '{field}' is required")
        
        # Create purchase record
        purchase = InsurancePurchase(
            user_id=request['user_id'],
            package_name=request['package_name'],
            package_type=request['package_type'],
            insurance_company=request.get('insurance_company'),
            customer_name=request['customer_name'],
            customer_phone=request['customer_phone'],
            customer_email=request.get('customer_email'),
            customer_address=request.get('customer_address'),
            customer_id_number=request.get('customer_id_number'),
            coverage_amount=request.get('coverage_amount'),
            premium_amount=request['premium_amount'],
            payment_frequency=request.get('payment_frequency'),
            start_date=request.get('start_date'),
            end_date=request.get('end_date'),
            beneficiary_name=request.get('beneficiary_name'),
            beneficiary_relationship=request.get('beneficiary_relationship'),
            vehicle_type=request.get('vehicle_type'),
            license_plate=request.get('license_plate'),
            payment_method=request.get('payment_method'),
            payment_status=request.get('payment_status', 'PENDING'),
            transaction_id=request.get('transaction_id'),
            document_id=request.get('document_id'),
            policy_number=request.get('policy_number'),
            status=request.get('status', 'ACTIVE'),
            additional_data=json.dumps(request.get('additional_data', {})) if request.get('additional_data') else None,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        db.add(purchase)
        db.commit()
        db.refresh(purchase)
        
        print(f"✅ Created insurance purchase #{purchase.id} for user {purchase.user_id}")
        
        return {
            "purchase_id": purchase.id,
            "message": "Insurance purchase created successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Failed to create purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to create purchase: {str(e)}")
    finally:
        db.close()


@app.get("/users/{user_id}/insurance-purchases")
async def get_user_insurance_purchases(user_id: int):
    """
    Get all insurance purchases for a user
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchases = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.user_id == user_id)\
            .order_by(InsurancePurchase.created_at.desc())\
            .all()
        
        result = []
        for p in purchases:
            result.append({
                "id": p.id,
                "package_name": p.package_name,
                "package_type": p.package_type,
                "insurance_company": p.insurance_company,
                "customer_name": p.customer_name,
                "customer_phone": p.customer_phone,
                "customer_email": p.customer_email,
                "customer_address": p.customer_address,
                "customer_id_number": p.customer_id_number,
                "coverage_amount": p.coverage_amount,
                "premium_amount": p.premium_amount,
                "payment_frequency": p.payment_frequency,
                "start_date": p.start_date,
                "end_date": p.end_date,
                "beneficiary_name": p.beneficiary_name,
                "beneficiary_relationship": p.beneficiary_relationship,
                "vehicle_type": p.vehicle_type,
                "license_plate": p.license_plate,
                "payment_method": p.payment_method,
                "payment_status": p.payment_status,
                "transaction_id": p.transaction_id,
                "document_id": p.document_id,
                "policy_number": p.policy_number,
                "status": p.status,
                "additional_data": json.loads(p.additional_data) if p.additional_data else None,
                "created_at": p.created_at.isoformat() if p.created_at else None,
                "updated_at": p.updated_at.isoformat() if p.updated_at else None
            })
        
        print(f"📋 Retrieved {len(result)} purchases for user {user_id}")
        
        return {
            "purchases": result,
            "total": len(result)
        }
        
    except Exception as e:
        print(f"❌ Failed to get purchases: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get purchases: {str(e)}")
    finally:
        db.close()


@app.get("/insurance-purchases/{purchase_id}")
async def get_insurance_purchase(purchase_id: int):
    """
    Get single insurance purchase by ID
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchase = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.id == purchase_id)\
            .first()
        
        if not purchase:
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        return {
            "id": purchase.id,
            "user_id": purchase.user_id,
            "package_name": purchase.package_name,
            "package_type": purchase.package_type,
            "insurance_company": purchase.insurance_company,
            "customer_name": purchase.customer_name,
            "customer_phone": purchase.customer_phone,
            "customer_email": purchase.customer_email,
            "customer_address": purchase.customer_address,
            "customer_id_number": purchase.customer_id_number,
            "coverage_amount": purchase.coverage_amount,
            "premium_amount": purchase.premium_amount,
            "payment_frequency": purchase.payment_frequency,
            "start_date": purchase.start_date,
            "end_date": purchase.end_date,
            "beneficiary_name": purchase.beneficiary_name,
            "beneficiary_relationship": purchase.beneficiary_relationship,
            "vehicle_type": purchase.vehicle_type,
            "license_plate": purchase.license_plate,
            "payment_method": purchase.payment_method,
            "payment_status": purchase.payment_status,
            "transaction_id": purchase.transaction_id,
            "document_id": purchase.document_id,
            "policy_number": purchase.policy_number,
            "status": purchase.status,
            "additional_data": json.loads(purchase.additional_data) if purchase.additional_data else None,
            "created_at": purchase.created_at.isoformat() if purchase.created_at else None,
            "updated_at": purchase.updated_at.isoformat() if purchase.updated_at else None
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Failed to get purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get purchase: {str(e)}")
    finally:
        db.close()


@app.put("/insurance-purchases/{purchase_id}")
async def update_insurance_purchase(purchase_id: int, request: dict):
    """
    Update insurance purchase (e.g., payment status, policy number)
    """
    from app.models import InsurancePurchase
    db = get_db()
    
    try:
        purchase = db.query(InsurancePurchase)\
            .filter(InsurancePurchase.id == purchase_id)\
            .first()
        
        if not purchase:
            raise HTTPException(status_code=404, detail="Purchase not found")
        
        # Update allowed fields
        updatable_fields = [
            'payment_status', 'transaction_id', 'policy_number', 'status',
            'payment_method', 'start_date', 'end_date', 'additional_data'
        ]
        
        for field in updatable_fields:
            if field in request:
                if field == 'additional_data' and request[field]:
                    setattr(purchase, field, json.dumps(request[field]))
                else:
                    setattr(purchase, field, request[field])
        
        purchase.updated_at = datetime.utcnow()
        
        db.commit()
        db.refresh(purchase)
        
        print(f"✅ Updated insurance purchase #{purchase_id}")
        
        return {
            "purchase_id": purchase.id,
            "message": "Insurance purchase updated successfully"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        db.rollback()
        print(f"❌ Failed to update purchase: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update purchase: {str(e)}")
    finally:
        db.close()


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)